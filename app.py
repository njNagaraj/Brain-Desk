from flask import Flask, request, jsonify, render_template, send_file
from werkzeug.utils import secure_filename
import openai
import os
import tempfile
from docx import Document
from pdfminer.high_level import extract_text
import io
import requests  # Added missing import
import json      # Added for OpenRouter API
from gtts import gTTS  # Google Text-to-Speech
import io

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 5 * 1024 * 1024  # 5MB limit

# Data model (simulating your DB)
AI_SOURCE = {
    'chatgpt': { 'name': 'chatgpt' },
    'deepseek': { 'name': 'deepseek' }
}

API_VERSIONS = {
    'gpt-4.0': { 'ai_source': 'chatgpt', 'name': 'gpt-4.0' },
    'deepseek/deepseek-r1-0528:free': { 'ai_source': 'deepseek', 'name': 'deepseek/deepseek-r1-0528:free' }
}

SETUP = {
    'api_key': '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',
    'ai_version': 'deepseek/deepseek-r1-0528:free'
}

PROCESSES = {
    'meeting_minutes_maker': {
        'process_name': 'meeting_minutes_maker',
        'ai_source': 'deepseek',
        'ai_version': 'deepseek/deepseek-r1-0528:free',
        'prompt_order': ['system']
    },
    'formal_email_genrator': {
        'process_name': 'formal_email_genrator',
        'ai_source': 'deepseek',
        'ai_version': 'deepseek/deepseek-r1-0528:free',
        'prompt_order': ['system']
    }
}

PROMPTS = [
    {
        'process': 'meeting_minutes_maker',
        'role': 'system',
        'prompt': 'You are meeting minutes maker based on the provided content. Always give a short 2 minutes meeting overview.'
    },
    {
        'process': 'formal_email_genrator',
        'role': 'system',
        'prompt': 'You are a formal email generator. Always generate a professional email based on the provided content. my name is Subash. CEO of Doumein, Contact: 6382476871'
    }
]

# Setup OpenAI API key (used only for chatgpt)
openai.api_key = SETUP['api_key']

@app.route('/', methods=['GET'])
def index():
    process_list = list(PROCESSES.keys())
    return render_template('index.html', processes=process_list)

@app.route('/ask', methods=['POST'])
def ask():
    print("\n===== New /ask request =====")
    
    try:
        # Get request parameters
        process_name = request.form.get('process')
        input_type = request.form.get('inputType')
        output_format = request.form.get('outputFormat')
        document_format = request.form.get('documentFormat', 'pdf')  # Default to PDF
        
        print(f"[DEBUG] Process: {process_name}, Input: {input_type}, Output: {output_format}, Doc Format: {document_format}")

        # Get content based on input type
        if input_type == 'text':
            content = request.form.get('content', '').strip()
            if not content:
                print("[ERROR] No text content provided")
                return jsonify({'error': 'Please enter a query before submitting.'}), 400
        else:
            if 'document' not in request.files:
                print("[ERROR] No file uploaded")
                return jsonify({'error': 'No file uploaded.'}), 400
                
            file = request.files['document']
            if file.filename == '':
                print("[ERROR] Empty file name")
                return jsonify({'error': 'No selected file.'}), 400
                
            filename = secure_filename(file.filename)
            ext = os.path.splitext(filename)[1].lower()
            print(f"[DEBUG] Processing file: {filename}, Extension: {ext}")
            
            # Save to temp file
            with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as temp:
                file.save(temp.name)
                temp_path = temp.name
            
            try:
                if ext == '.pdf':
                    print("[DEBUG] Extracting text from PDF")
                    content = extract_text(temp_path)
                elif ext in ('.doc', '.docx'):
                    print("[DEBUG] Extracting text from Word document")
                    doc = Document(temp_path)
                    content = '\n'.join([para.text for para in doc.paragraphs])
                elif ext == '.txt':
                    print("[DEBUG] Reading text file")
                    with open(temp_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                else:
                    print(f"[ERROR] Unsupported file type: {ext}")
                    return jsonify({'error': 'Unsupported file type.'}), 400
            except Exception as e:
                print(f"[ERROR] File processing error: {str(e)}")
                return jsonify({'error': f'Error processing file: {str(e)}'}), 500
            finally:
                os.unlink(temp_path)
                print("[DEBUG] Temporary file deleted")

        print(f"[DEBUG] Content length: {len(content)} characters")
        
        # Process lookup and message building
        process = PROCESSES.get(process_name)
        if not process:
            print(f"[ERROR] Invalid process selected: {process_name}")
            return jsonify({'error': 'Invalid process selected.'}), 400

        ai_source = process['ai_source']
        ai_version = process['ai_version']
        prompt_order = process['prompt_order']

        print(f"[DEBUG] AI Source: {ai_source}, Version: {ai_version}")
        
        # Build messages
        messages = []
        for role in prompt_order:
            matching_prompt = next((p for p in PROMPTS if p['process'] == process_name and p['role'] == role), None)
            if matching_prompt:
                messages.append({'role': role, 'content': matching_prompt['prompt']})
                print(f"[DEBUG] Added {role} prompt")

        messages.append({'role': 'user', 'content': content})
        print("[DEBUG] Messages prepared for AI")
        
        # Call AI API
        try:
            if ai_source == 'chatgpt':
                print("[DEBUG] Calling OpenAI API")
                response = openai.ChatCompletion.create(
                    model=ai_version,
                    messages=messages
                )
                reply = response['choices'][0]['message']['content'].strip()
                print("[DEBUG] Received response from OpenAI")

            elif ai_source == 'deepseek':
                print("[DEBUG] Calling OpenRouter API for DeepSeek")
                correction_response = requests.post(
                    url="https://openrouter.ai/api/v1/chat/completions",
                    headers={
                        "Authorization": f"Bearer {SETUP['api_key']}",
                        "Content-Type": "application/json"
                    },
                    data=json.dumps({
                        "model": ai_version,
                        "messages": messages,
                        "temperature": 0.1
                    })
                )
                print(f"[DEBUG] OpenRouter status: {correction_response.status_code}")
                reply = correction_response.json()['choices'][0]['message']['content'].strip()
                print("[DEBUG] Received response from DeepSeek")
            else:
                print("[ERROR] Unsupported AI source")
                return jsonify({'error': 'Unsupported AI source.'}), 400
        except Exception as e:
            print(f"[ERROR] AI API call failed: {str(e)}")
            return jsonify({'error': f'Error calling AI: {str(e)}'}), 500

        print(f"[DEBUG] AI response length: {len(reply)} characters")
        
        # Handle output format
        if output_format == 'text':
            print("[DEBUG] Returning text response")
            return reply
        elif output_format == 'voice':
            print("[DEBUG] Generating voice output")
            try:
                tts = gTTS(text=reply, lang='en')
                audio_file = io.BytesIO()
                tts.write_to_fp(audio_file)
                audio_file.seek(0)
                return send_file(
                    audio_file,
                    mimetype='audio/mpeg',
                    as_attachment=True,
                    download_name='output.mp3'
                )
            except Exception as e:
                print(f"[ERROR] Voice generation failed: {str(e)}")
                return jsonify({'error': f'Error generating voice: {str(e)}'}), 500
        elif output_format == 'document':
            print(f"[DEBUG] Generating {document_format} document")
            try:
                output = io.BytesIO()
                
                if document_format == 'pdf':
                    from reportlab.pdfgen import canvas
                    from reportlab.lib.pagesizes import letter
                    
                    p = canvas.Canvas(output, pagesize=letter)
                    p.setFont("Helvetica", 12)
                    
                    # Simple text wrapping
                    lines = []
                    words = reply.split()
                    line = ''
                    for word in words:
                        if len(line) + len(word) < 80:
                            line += word + ' '
                        else:
                            lines.append(line)
                            line = word + ' '
                    if line:
                        lines.append(line)
                    
                    # Add text to PDF
                    y = 750
                    for line in lines:
                        p.drawString(50, y, line)
                        y -= 15
                        if y < 50:
                            p.showPage()
                            y = 750
                    p.save()
                    mimetype = 'application/pdf'
                    ext = 'pdf'
                else:  # doc
                    doc = Document()
                    doc.add_paragraph(reply)
                    doc.save(output)
                    mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    ext = 'docx'
                
                output.seek(0)
                return send_file(
                    output,
                    mimetype=mimetype,
                    as_attachment=True,
                    download_name=f'output.{ext}'
                )
            except Exception as e:
                print(f"[ERROR] Document generation failed: {str(e)}")
                return jsonify({'error': f'Error generating document: {str(e)}'}), 500
        else:
            print(f"[ERROR] Unsupported output format: {output_format}")
            return jsonify({'error': 'Unsupported output format.'}), 400

    except Exception as e:
        print(f"[CRITICAL] Unhandled exception: {str(e)}")
        return jsonify({'error': 'An unexpected error occurred.'}), 500

if __name__ == '__main__':
    app.run(debug=True)